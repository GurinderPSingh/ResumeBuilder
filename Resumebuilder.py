
import sys
print(sys.executable)
from docx import Document
# Create a Word document
doc = Document()

# Add content to the document
doc.add_heading("Gurinder Pal Singh", level=0)
doc.add_paragraph("(204) 430 2341 | gurinderpalsingh900@gmail.com")

doc.add_section()

doc.add_heading("Professional Summary", level=1)
doc.add_paragraph(
    "Experienced Data Engineer and Solution Specialist with over 7 years of expertise in designing and implementing data "
    "strategies. Proficient in cloud-based data ecosystems (Azure and GCP), advanced ETL development, and creating impactful "
    "data integrations. Strong leader adept at facilitating cross-functional collaboration, managing technical teams, and "
    "delivering data-driven solutions. Proven success in utilizing tools like Dayforce, Azure Data Factory, Power BI, and Python "
    "to streamline operations and drive business outcomes."
)

doc.add_section()

doc.add_heading("Technical Skills", level=1)
doc.add_paragraph(
    "- Statistical Tools: Alteryx, SAS, MATLAB\n"
    "- Cloud Platforms: Azure (Data Lake, Data Factory, Synapse, ML), GCP\n"
    "- ETL Tools: Azure Data Factory, Google BigQuery, Power Query\n"
    "- Programming & Analytics: Python, SQL, R, JSP\n"
    "- Data Visualization: Power BI, Tableau, Looker Studio\n"
    "- Integration & Architecture: REST APIs, Advanced Data Models, Data Schemas\n"
    "- Project Management Tools: Jira, Confluence, Agile/Scrum\n"
    "- Collaboration Tools: Jira, Asana, Slack, Microsoft Teams"
)

doc.add_section()

doc.add_heading("Professional Certifications", level=1)
doc.add_paragraph(
    "- Microsoft Certified: Azure Data Engineer Associate\n"
    "- Microsoft Certified: Power BI Data Analyst (PL-300)\n"
    "- Microsoft Certified: Azure Fundamentals\n"
    "- Certified Project Management Professional (PMP)"
)

doc.add_section()

doc.add_heading("Professional Experience", level=1)

doc.add_heading("Data Integration Specialist | SkipTheDishes", level=2)
doc.add_paragraph("Jan 2018 – Apr 2023")
doc.add_paragraph(
    "- Designed and maintained scalable ETL solutions, integrating Dayforce data with operational databases to optimize workforce management.\n"
    "- Implemented predictive models using Python and Azure ML, achieving 90% accuracy in demand forecasting.\n"
    "- Developed reusable Power BI templates and dashboards, enabling stakeholders to gain actionable insights in real time.\n"
    "- Enhanced operational efficiency by automating reporting processes, reducing manual effort by 60%.\n"
    "- Collaborated with IT and business teams in agile environments, delivering data solutions on time and within budget.\n"
    "- Integrated Git-based version control for ETL scripts and workflows, ensuring traceability and facilitating collaboration.\n"
    "- Migrated legacy data workflows to Google BigQuery (GBQ), enhancing performance and scalability.\n"
    "- Implemented Dayforce-specific analytics to optimize employee scheduling and payroll processes."
)

doc.add_heading("Senior Data Analyst | SkipTheDishes", level=2)
doc.add_paragraph("Sep 2019 – Apr 2023")
doc.add_paragraph(
    "- Created interactive dashboards in Power BI for real-time operational and strategic insights.\n"
    "- Improved reporting accuracy by automating manual processes using Python and SQL.\n"
    "- Supported leadership with data-driven insights, resulting in a 22% cost reduction through targeted operational improvements.\n"
    "- Mentored a team of junior analysts, fostering a culture of collaboration and knowledge sharing.\n"
    "- Designed advanced integrations between Dayforce and operational databases, enabling streamlined workforce management.\n"
    "- Implemented Git-based workflows for version control of data visualization assets.\n"
    "- Leveraged Google BigQuery (GBQ) to centralize and optimize data analytics processes, reducing query execution times by 35%."
)

doc.add_heading("Senior Data Engineer | Red River College", level=2)
doc.add_paragraph("May 2023 – Present")
doc.add_paragraph(
    "- Automated ETL workflows and integrated systems using Azure Data Factory and Power BI, reducing data processing time by 40%.\n"
    "- Designed advanced data schemas and models for optimized data storage and retrieval, ensuring seamless scalability and compliance.\n"
    "- Spearheaded resource allocation models, reducing program waitlists by 30%.\n"
    "- Partnered with leadership to implement robust data quality benchmarks and improve decision-making frameworks.\n"
    "- Acted as a technical lead for cross-departmental data integration projects, ensuring alignment with business objectives.\n"
    "- Developed Python-based automation scripts to streamline data validation and preprocessing, significantly reducing manual intervention."
)

doc.add_section()

doc.add_heading("Key Projects", level=1)
doc.add_paragraph(
    "- Dayforce Data Integration**: Automated data extraction and integration processes between Dayforce and Azure Data Factory, ensuring accurate and timely reporting.\n"
    "- Cloud Migration & Optimization**: Migrated legacy systems to Azure, achieving a 25% reduction in infrastructure costs and enhancing data scalability.\n"
    "- Predictive Analytics**: Developed demand forecasting models using Python and Azure ML, enabling efficient resource planning during peak periods.\n"
    "- Centralized Data Catalog**: Implemented a searchable catalog for metadata management, reducing data retrieval time by 30%.\n"
    "- Enterprise Governance Framework**: Established data governance protocols for compliance and consistency across systems."
)

doc.add_section()

doc.add_heading("Education", level=1)
doc.add_paragraph(
    "- Professional Development Certificate in Data Science and Machine Learning (Ongoing)**\n  McGill University\n\n"
    "- Master’s in Data Sciences**\n  Carleton University\n\n"
    "- Master’s in Engineering (Electronics)**\n  Punjab Technical University\n\n"
    "- Bachelor’s in Engineering (Electronics)**\n  Punjab Technical University"
)

# Save the document
file_path = "C:/ResumeBuilder/Tailored_Resume_Gurinder.docx"
doc.save(file_path)

file_path
